from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[3]
SOURCE = ROOT / "deliverables" / "YouTube财经博主内容结构与钩子仿写手册_Ben完整交付版.md"
OUTPUT = ROOT / "deliverables" / "YouTube财经博主内容结构与钩子仿写手册_Ben完整交付版.docx"

BLUE = "183B56"
TEAL = "0E7490"
PALE_BLUE = "EAF2F6"
PALE_TEAL = "E8F5F5"
PALE_GOLD = "FFF7E6"
GRAY = "5F6B76"
LIGHT_GRAY = "F3F5F7"
WHITE = "FFFFFF"
EAST_ASIA_FONT = "Microsoft YaHei"


def set_run_font(run, *, size=None, bold=None, color=None, italic=None, name="Calibri"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instr, separate, text, end):
        run._r.append(node)
    set_run_font(run, size=9, color=GRAY)


def create_numbering_instance(doc):
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=-1) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi_level = OxmlElement("w:multiLevelType")
    multi_level.set(qn("w:val"), "singleLevel")
    abstract.append(multi_level)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "%1.")
    suffix = OxmlElement("w:suff")
    suffix.set(qn("w:val"), "tab")
    level.append(start)
    level.append(num_fmt)
    level.append(level_text)
    level.append(suffix)
    abstract.append(level)
    first_num = numbering.find(qn("w:num"))
    if first_num is None:
        numbering.append(abstract)
    else:
        first_num.addprevious(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    else:
        for child in list(num_pr):
            num_pr.remove(child)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_node = OxmlElement("w:numId")
    num_id_node.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_node)


def add_inline(paragraph, text, *, base_size=11, base_color="263238"):
    token_re = re.compile(r"(\*\*.*?\*\*|`.*?`)")
    cursor = 0
    for match in token_re.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor : match.start()])
            set_run_font(run, size=base_size, color=base_color)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=base_size, bold=True, color=BLUE)
        else:
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, size=base_size, bold=True, color=TEAL)
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        set_run_font(run, size=base_size, color=base_color)


def add_callout(doc, text, fill=PALE_BLUE, border=TEAL):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.22
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    p_bdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), border)
    p_bdr.append(left)
    p_pr.append(p_bdr)
    add_inline(p, text, base_size=10.5)
    return p


def setup_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string("263238")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, TEAL, 14, 7),
        ("Heading 3", 11.5, BLUE, 10, 5),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    styles["Heading 1"].paragraph_format.page_break_before = True

    for name, left, hanging in (("List Bullet", 0.38, 0.19), ("List Number", 0.38, 0.19)):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(left)
        style.paragraph_format.first_line_indent = Inches(-hanging)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    if "TOC Entry" not in styles:
        toc = styles.add_style("TOC Entry", WD_STYLE_TYPE.PARAGRAPH)
    else:
        toc = styles["TOC Entry"]
    toc.font.name = "Calibri"
    toc._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    toc.font.size = Pt(11)
    toc.font.color.rgb = RGBColor.from_string(BLUE)
    toc.paragraph_format.space_after = Pt(5)
    toc.paragraph_format.left_indent = Inches(0.2)


def add_cover(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.49)
    section.footer_distance = Inches(0.49)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(54)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("YouTube 财经博主")
    set_run_font(run, size=28, bold=True, color=BLUE)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    run = p.add_run("内容结构与钩子仿写手册")
    set_run_font(run, size=25, bold=True, color=TEAL)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(28)
    run = p.add_run("给 Ben 哥的完整交付版")
    set_run_font(run, size=13, color=GRAY)

    add_callout(
        doc,
        "看到一个财经选题后，怎样把它写成观众愿意点开、能听下去、最后知道该观察什么的视频。",
        fill=PALE_TEAL,
        border=TEAL,
    )

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(28)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("三种核心内容任务")
    set_run_font(run, size=12, bold=True, color=BLUE)

    for label, detail in (
        ("决策型", "今天这件事到底怎么看"),
        ("讲懂型", "这个复杂概念到底是什么"),
        ("解释型", "宏观变化为什么会影响产业和资产"),
    ):
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(f"{label}：")
        set_run_font(r, size=11, bold=True, color=TEAL)
        r = p.add_run(detail)
        set_run_font(r, size=11, color="263238")

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(34)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("内容训练手册  |  2026年8月")
    set_run_font(r, size=9.5, color=GRAY)

    doc.add_page_break()


def add_toc(doc, lines):
    p = doc.add_paragraph()
    r = p.add_run("阅读导航")
    set_run_font(r, size=18, bold=True, color=BLUE)
    p.paragraph_format.space_after = Pt(14)

    add_callout(doc, "先读第一部分掌握总原则；需要写稿时，直接跳到第七部分选择模板；审稿时使用第十部分检查表。", fill=PALE_GOLD, border="C58A00")

    h1s = [line[2:].strip() for line in lines if line.startswith("# ")][1:]
    for text in h1s:
        p = doc.add_paragraph(style="TOC Entry")
        add_inline(p, text, base_size=11, base_color=BLUE)
    doc.add_page_break()


def parse_table(lines, start):
    rows = []
    idx = start
    while idx < len(lines) and lines[idx].startswith("|"):
        cells = [cell.strip() for cell in lines[idx].strip().strip("|").split("|")]
        rows.append(cells)
        idx += 1
    if len(rows) >= 2 and all(re.fullmatch(r":?-{3,}:?", cell) for cell in rows[1]):
        rows.pop(1)
    return rows, idx


def add_table(doc, rows):
    cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    widths = [1580, 3180, 1960, 2640] if cols == 4 else [9360 // cols] * cols
    set_table_geometry(table, widths)
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if r_idx == 0:
                set_cell_shading(cell, BLUE)
            elif r_idx % 2 == 0:
                set_cell_shading(cell, LIGHT_GRAY)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.12
            add_inline(p, value, base_size=9.2, base_color=WHITE if r_idx == 0 else "263238")
            for run in p.runs:
                if r_idx == 0:
                    run.bold = True
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_body(doc, lines):
    idx = 0
    seen_first_h1 = False
    body_h1_count = 0
    current_num_id = None
    flow_h1s = {
        "九、从热点到成稿的工作顺序",
        "十、Ben 哥审稿检查表",
        "十一、一页速查卡",
        "十二、这版手册目前能怎么用",
    }
    while idx < len(lines):
        raw = lines[idx].rstrip()
        text = raw.strip()
        if not text or text == "---":
            idx += 1
            continue
        if raw.startswith("# "):
            current_num_id = None
            if not seen_first_h1:
                seen_first_h1 = True
                idx += 1
                continue
            heading_text = raw[2:].strip()
            p = doc.add_paragraph(style="Heading 1")
            if body_h1_count == 0 or heading_text in flow_h1s:
                p.paragraph_format.page_break_before = False
            body_h1_count += 1
            add_inline(p, heading_text, base_size=16, base_color=BLUE)
            idx += 1
            continue
        if raw.startswith("## "):
            current_num_id = None
            p = doc.add_paragraph(style="Heading 2")
            add_inline(p, raw[3:].strip(), base_size=13, base_color=TEAL)
            idx += 1
            continue
        if raw.startswith("### "):
            current_num_id = None
            p = doc.add_paragraph(style="Heading 3")
            add_inline(p, raw[4:].strip(), base_size=11.5, base_color=BLUE)
            idx += 1
            continue
        if raw.startswith("|"):
            current_num_id = None
            rows, idx = parse_table(lines, idx)
            add_table(doc, rows)
            continue
        if raw.startswith("> "):
            current_num_id = None
            block = [raw[2:].strip()]
            idx += 1
            while idx < len(lines) and lines[idx].startswith("> "):
                block.append(lines[idx][2:].strip())
                idx += 1
            add_callout(doc, " ".join(block), fill=PALE_BLUE, border=TEAL)
            continue
        if re.match(r"^- ", raw):
            current_num_id = None
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, raw[2:].strip())
            idx += 1
            continue
        if re.match(r"^\d+\. ", raw):
            if current_num_id is None:
                current_num_id = create_numbering_instance(doc)
            p = doc.add_paragraph(style="List Number")
            apply_numbering(p, current_num_id)
            add_inline(p, re.sub(r"^\d+\. ", "", raw))
            idx += 1
            continue
        if text.startswith("`") and text.endswith("`"):
            current_num_id = None
            add_callout(doc, text[1:-1], fill=PALE_TEAL, border=TEAL)
            idx += 1
            continue
        current_num_id = None
        p = doc.add_paragraph()
        add_inline(p, text)
        idx += 1


def add_running_furniture(doc):
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(0.85)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.header_distance = Inches(0.42)
        section.footer_distance = Inches(0.42)
        section.different_first_page_header_footer = True

        header = section.header
        p = header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = p.add_run("YouTube 财经内容仿写手册")
        set_run_font(r, size=8.5, color=GRAY)

        footer = section.footer
        p = footer.paragraphs[0]
        add_page_number(p)


def main():
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    doc = Document()
    setup_styles(doc)
    add_cover(doc)
    add_toc(doc, lines)
    add_body(doc, lines)
    add_running_furniture(doc)

    doc.core_properties.title = "YouTube 财经博主内容结构与钩子仿写手册"
    doc.core_properties.subject = "给 Ben 哥的可直接执行内容训练手册"
    doc.core_properties.author = "Content Research Team"
    doc.core_properties.keywords = "YouTube, 财经内容, 钩子, 叙事结构, 脚本模板"
    doc.core_properties.comments = "学习结构，不复制人格与原句。"

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
